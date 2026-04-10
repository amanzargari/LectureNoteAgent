from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

from .io_utils import SlideImageAsset


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_IMAGE_RE = re.compile(r"!\[(?P<caption>[^\]]*)\]\((?P<target>[^)]+)\)")
_TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


def _clean_inline_markdown(text: str) -> str:
    out = text
    out = re.sub(r"`([^`]+)`", r"\1", out)
    out = re.sub(r"\*\*([^*]+)\*\*", r"\1", out)
    out = re.sub(r"__([^_]+)__", r"\1", out)
    out = re.sub(r"\*([^*]+)\*", r"\1", out)
    out = re.sub(r"_([^_]+)_", r"\1", out)
    out = out.replace("\\[", "[").replace("\\]", "]")
    return out.strip()


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _is_markdown_table(lines: list[str], start_idx: int) -> bool:
    if start_idx + 1 >= len(lines):
        return False
    first = lines[start_idx].strip()
    second = lines[start_idx + 1].strip()
    return first.startswith("|") and bool(_TABLE_SEP_RE.match(second))


def _build_image_index(slide_images: dict[int, list[SlideImageAsset]]) -> dict[str, SlideImageAsset]:
    index: dict[str, SlideImageAsset] = {}
    for assets in slide_images.values():
        for asset in assets:
            key = (asset.image_ref or "").strip().lower()
            if key and key not in index:
                index[key] = asset
    return index


def _normalize_image_ref(target: str) -> str:
    t = (target or "").strip()
    for prefix in ("image_ref:", "ref:", "image:"):
        if t.lower().startswith(prefix):
            t = t[len(prefix) :].strip()
            break
    return t.strip().lower()


def _tokenize_ref(text: str) -> set[str]:
    return {tok for tok in re.split(r"[^a-z0-9]+", (text or "").lower()) if tok}


def _resolve_image_asset(ref_key: str, image_index: dict[str, SlideImageAsset]) -> SlideImageAsset | None:
    if not ref_key:
        return None

    # 1) Exact match.
    direct = image_index.get(ref_key)
    if direct is not None:
        return direct

    # 2) Match by stem/path-like reference.
    stem = Path(ref_key).stem.lower()
    for key, asset in image_index.items():
        if key == stem or Path(key).stem.lower() == stem:
            return asset

    # 3) Substring match (common when model emits variants/extensions).
    for key, asset in image_index.items():
        if ref_key in key or key in ref_key:
            return asset

    # 4) Token-overlap fallback.
    target_tokens = _tokenize_ref(ref_key)
    best_asset: SlideImageAsset | None = None
    best_score = 0.0
    if target_tokens:
        for key, asset in image_index.items():
            key_tokens = _tokenize_ref(key)
            if not key_tokens:
                continue
            overlap = len(target_tokens & key_tokens)
            if overlap == 0:
                continue
            score = overlap / max(1, len(target_tokens | key_tokens))
            if score > best_score:
                best_score = score
                best_asset = asset

    if best_score >= 0.5:
        return best_asset

    return None


def _normalize_figure_label(label: str) -> str:
    cleaned = _clean_inline_markdown(label or "")
    if not cleaned:
        return ""
    cleaned = re.sub(r"(?i)^\s*(?:figure|fig\.?)(?:\s*\d+)?\s*[:.)-]\s*", "", cleaned)
    return cleaned.strip()


def _compute_dynamic_image_width(img_path: Path, caption: str, context_line: str) -> float:
    max_w = 6.3
    min_w = 2.2
    signal = f"{caption} {context_line}".lower()

    try:
        from PIL import Image

        with Image.open(img_path) as img:
            width_px, height_px = img.size
    except Exception:
        return 5.4

    if height_px <= 0:
        return 5.4

    ratio = width_px / max(height_px, 1)
    natural = max(min_w, min(max_w, width_px / 220.0))

    if ratio > 1.9:
        base = 6.2
    elif ratio > 1.35:
        base = 5.8
    elif ratio < 0.8:
        base = 4.2
    else:
        base = 5.0

    if any(k in signal for k in ("icon", "logo", "small", "thumbnail")):
        base = min(base, 3.0)
    if any(k in signal for k in ("diagram", "architecture", "workflow", "pipeline", "chart", "table")):
        base = max(base, 5.8)

    width_in = (base + natural) / 2.0
    return max(min_w, min(max_w, width_in))


def _add_caption(document: Document, text: str) -> None:
    try:
        document.add_paragraph(text, style="Caption")
    except Exception:
        document.add_paragraph(text)


def _add_picture_with_fallback(run, img_path: Path, fallback_dir: Path, width_inches: float) -> bool:
    try:
        run.add_picture(str(img_path), width=Inches(width_inches))
        return True
    except Exception:
        try:
            from PIL import Image

            with Image.open(img_path) as img:
                converted = img.convert("RGB") if img.mode in {"RGBA", "P", "LA"} else img
                fallback_path = fallback_dir / f"{img_path.stem}_docx_safe.png"
                converted.save(fallback_path, format="PNG")

            run.add_picture(str(fallback_path), width=Inches(width_inches))
            return True
        except Exception:
            return False


def _add_markdown_body(document: Document, markdown_text: str, slide_images: dict[int, list[SlideImageAsset]]) -> None:
    lines = markdown_text.splitlines()
    image_index = _build_image_index(slide_images)
    figure_counter = 0
    table_counter = 0

    with tempfile.TemporaryDirectory(prefix="docx_img_fallback_") as tmp_dir:
        fallback_dir = Path(tmp_dir)
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            if not line.strip():
                document.add_paragraph("")
                i += 1
                continue

            # Markdown table block -> DOCX table
            if _is_markdown_table(lines, i):
                header_cells = _split_table_row(lines[i])
                rows: list[list[str]] = [header_cells]
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    rows.append(_split_table_row(lines[j]))
                    j += 1

                max_cols = max((len(r) for r in rows), default=1)
                table = document.add_table(rows=len(rows), cols=max_cols)
                table.style = "Table Grid"
                for r_idx, row_cells in enumerate(rows):
                    for c_idx in range(max_cols):
                        txt = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        table.cell(r_idx, c_idx).text = _clean_inline_markdown(txt)

                table_counter += 1
                _add_caption(document, f"Table {table_counter}")
                i = j
                continue

            # Inline image markers in markdown syntax
            image_matches = list(_IMAGE_RE.finditer(line))
            if image_matches:
                last = 0
                for m in image_matches:
                    prefix = line[last : m.start()].strip()
                    if prefix:
                        document.add_paragraph(_clean_inline_markdown(prefix))

                    caption = _clean_inline_markdown(m.group("caption") or "")
                    ref_key = _normalize_image_ref(m.group("target") or "")
                    asset = _resolve_image_asset(ref_key, image_index)

                    if asset is not None:
                        img_path = Path(asset.image_path)
                        if img_path.exists():
                            figure_counter += 1
                            width_in = _compute_dynamic_image_width(img_path, caption, line)
                            pic_paragraph = document.add_paragraph()
                            run = pic_paragraph.add_run()
                            ok = _add_picture_with_fallback(run, img_path, fallback_dir, width_in)
                            if ok:
                                figure_label = _normalize_figure_label(caption) or asset.image_ref
                                _add_caption(document, f"Figure {figure_counter}: {figure_label}")
                            else:
                                document.add_paragraph(f"(Skipped image: {asset.image_ref})")

                    last = m.end()

                suffix = line[last:].strip()
                if suffix:
                    document.add_paragraph(_clean_inline_markdown(suffix))
                i += 1
                continue

            heading_match = _HEADING_RE.match(line)
            if heading_match:
                level = min(len(heading_match.group(1)), 6)
                text = _clean_inline_markdown(heading_match.group(2))
                document.add_heading(text, level=level)
                i += 1
                continue

            bullet_match = _BULLET_RE.match(line)
            if bullet_match:
                document.add_paragraph(_clean_inline_markdown(bullet_match.group(1)), style="List Bullet")
                i += 1
                continue

            numbered_match = _NUMBERED_RE.match(line)
            if numbered_match:
                document.add_paragraph(_clean_inline_markdown(numbered_match.group(1)), style="List Number")
                i += 1
                continue

            document.add_paragraph(_clean_inline_markdown(line))
            i += 1


def write_docx_from_markdown(
    *,
    markdown_text: str,
    output_path: str,
    course_name: str,
    slide_images: dict[int, list[SlideImageAsset]] | None = None,
) -> None:
    doc = Document()
    title = (course_name or "Lecture Notes").strip()
    doc.add_heading(title, level=0)

    _add_markdown_body(doc, markdown_text, slide_images or {})

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
